"""
Output Formatter
Combines draft + all reports into a single structured output.
Saves to /outputs/ as Markdown, JSON, and optionally DOCX.
"""

import os
import json
from datetime import datetime


def format_final_output(
      topic: str,
      doc_type: str,
      draft: str,
      fact_report: dict,
      hallucination_report: dict,
      ai_detection_report: dict,
      plagiarism_report: dict
) -> dict:
      """Format and save the complete pipeline output."""

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_topic = "".join(c if c.isalnum() or c in "-_" else "_" for c in topic[:40])
    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)
    base_filename = f"{output_dir}/{timestamp}_{safe_topic}"

    # Compute overall quality score
    fact_score = fact_report.get("fact_score", 100)
    halluc_score = (1 - hallucination_report.get("overall_score", 0.5)) * 100
    ai_score = (1 - ai_detection_report.get("overall_ai_probability", 0.5)) * 100
    plagiarism_score = (1 - plagiarism_report.get("overall_similarity", 0.0)) * 100
    overall_quality = round((fact_score + ai_score + plagiarism_score) / 3, 1)

    output = {
              "metadata": {
                            "topic": topic,
                            "doc_type": doc_type,
                            "generated_at": timestamp,
                            "word_count": len(draft.split()),
                            "overall_quality_score": overall_quality
              },
              "draft": draft,
              "reports": {
                            "fact_check": fact_report,
                            "hallucination": hallucination_report,
                            "ai_detection": ai_detection_report,
                            "plagiarism": plagiarism_report
              }
    }

    # Save JSON (full data)
    json_path = f"{base_filename}.json"
    with open(json_path, "w", encoding="utf-8") as f:
              json.dump(output, f, indent=2, ensure_ascii=False)

    # Save Markdown (human-readable)
    md_path = f"{base_filename}.md"
    md_content = _build_markdown(topic, doc_type, draft, fact_report,
                                                                   hallucination_report, ai_detection_report,
                                                                   plagiarism_report, overall_quality, timestamp)
    with open(md_path, "w", encoding="utf-8") as f:
              f.write(md_content)

    # Try to save DOCX
    try:
              _save_docx(draft, topic, doc_type, f"{base_filename}.docx")
              output["docx_path"] = f"{base_filename}.docx"
              print(f"  Saved DOCX: {base_filename}.docx")
except ImportError:
          print("  (python-docx not installed - skipping DOCX export)")

    print(f"  Saved JSON: {json_path}")
    print(f"  Saved Markdown: {md_path}")
    print(f"  Overall Quality Score: {overall_quality}/100")

    return output


def _build_markdown(topic, doc_type, draft, fact_report, halluc_report,
                                        ai_report, plag_report, quality_score, timestamp) -> str:
                                              return f"""# Smart Writer Stack - Output Report
                                              **Topic:** {topic}
                                              **Document Type:** {doc_type.upper()}
                                              **Generated:** {timestamp}
                                              **Overall Quality Score:** {quality_score}/100

                                              ---

                                              ## Quality Scorecard

                                              | Check | Result | Score |
                                              |-------|--------|-------|
                                              | Fact Check | {fact_report.get('supported', 0)}/{fact_report.get('total_claims', 0)} claims supported | {fact_report.get('fact_score', 'N/A')}% |
                                              | Hallucination Risk | {halluc_report.get('risk_level', 'N/A')} | {halluc_report.get('overall_score', 0):.3f} similarity |
                                              | AI Detection | {ai_report.get('detection_verdict', 'N/A')} | {ai_report.get('overall_ai_probability', 0):.1%} AI |
                                              | Plagiarism | {plag_report.get('plagiarism_verdict', 'N/A')} | {plag_report.get('overall_similarity', 0):.1%} similar |

                                              ---

                                              ## Generated Document

                                              {draft}

                                              ---

                                              ## Fact Check Details

                                              - Total claims checked: {fact_report.get('total_claims', 0)}
                                              - Supported: {fact_report.get('supported', 0)}
                                              - Unsupported: {fact_report.get('unsupported', 0)}
                                              - Partial: {fact_report.get('partial', 0)}

                                              ---

                                              ## Flagged Sentences (Possible Hallucinations)

                                              {chr(10).join(f"- [{r['risk']}] {r['sentence'][:100]}..." for r in halluc_report.get('flagged_sentences', [])[:5])}

                                              ---
                                              *Generated by Smart Writer Stack | github.com/ekawough/smart-writer-stack*
                                              """


def _save_docx(draft: str, topic: str, doc_type: str, path: str):
      from docx import Document
      doc = Document()
      doc.add_heading(topic, 0)
      doc.add_paragraph(f"Document Type: {doc_type.upper()}")
      doc.add_paragraph("")
      for paragraph in draft.split("\n\n"):
                if paragraph.strip():
                              if paragraph.startswith("#"):
                                                level = len(paragraph) - len(paragraph.lstrip("#"))
                                                doc.add_heading(paragraph.lstrip("# "), min(level, 4))
                else:
                                  doc.add_paragraph(paragraph.strip())
                      doc.save(path)
